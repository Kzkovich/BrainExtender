"""
Parse PDF, DOCX, and images into text + extracted image files.
"""
import base64
import io
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

MAX_IMAGES_PER_DOC = 10   # не извлекаем больше чем нужно
MAX_IMAGE_BYTES = 5 * 1024 * 1024  # 5MB на одно изображение


@dataclass
class ParsedDocument:
    text: str
    images: list[str] = field(default_factory=list)   # relative paths inside brain
    source_filename: str = ""
    page_count: int = 0
    error: Optional[str] = None


def _save_image_binary(storage, img_bytes: bytes, ext: str) -> Optional[str]:
    """Save raw image bytes to attachments/ folder. Returns relative path or None."""
    if len(img_bytes) > MAX_IMAGE_BYTES:
        return None  # слишком большое — пропускаем
    try:
        img_name = f"{uuid.uuid4().hex[:8]}.{ext}"
        img_rel_path = f"attachments/{img_name}"
        img_abs = storage.root / img_rel_path
        img_abs.parent.mkdir(parents=True, exist_ok=True)
        img_abs.write_bytes(img_bytes)
        return img_rel_path
    except Exception:
        return None


async def parse_pdf(file_bytes: bytes, filename: str, storage) -> ParsedDocument:
    """Extract text and images from PDF using pymupdf."""
    try:
        import fitz  # pymupdf

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        image_paths = []
        img_count = 0

        for page_num, page in enumerate(doc):
            # Text — strip null bytes and control chars that break the API
            text = page.get_text("text").strip()
            text = "".join(ch for ch in text if ch >= " " or ch in "\n\t")
            if text:
                text_parts.append(f"--- Страница {page_num + 1} ---\n{text}")

            # Images (limit total)
            if img_count < MAX_IMAGES_PER_DOC:
                for img in page.get_images(full=True):
                    if img_count >= MAX_IMAGES_PER_DOC:
                        break
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        img_bytes = base_image["image"]
                        img_ext = base_image.get("ext", "png")
                        path = _save_image_binary(storage, img_bytes, img_ext)
                        if path:
                            image_paths.append(path)
                            img_count += 1
                    except Exception:
                        continue

        page_count = len(doc)
        doc.close()

        return ParsedDocument(
            text="\n\n".join(text_parts),
            images=image_paths,
            source_filename=filename,
            page_count=page_count,
        )
    except ImportError:
        return ParsedDocument(
            text="",
            error="pymupdf не установлен на сервере.",
            source_filename=filename,
        )
    except Exception as e:
        return ParsedDocument(text="", error=f"Ошибка PDF: {e}", source_filename=filename)


async def parse_docx(file_bytes: bytes, filename: str, storage) -> ParsedDocument:
    """Extract text and images from DOCX."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        image_paths = []
        img_count = 0

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                text_parts.append("\n".join(rows))

        for rel in doc.part.rels.values():
            if img_count >= MAX_IMAGES_PER_DOC:
                break
            if "image" in rel.reltype:
                try:
                    img_bytes = rel.target_part.blob
                    content_type = rel.target_part.content_type
                    ext = content_type.split("/")[-1].split("+")[0]
                    if ext in ("jpeg", "jpg", "png", "gif", "webp"):
                        path = _save_image_binary(storage, img_bytes, ext)
                        if path:
                            image_paths.append(path)
                            img_count += 1
                except Exception:
                    continue

        return ParsedDocument(
            text="\n\n".join(text_parts),
            images=image_paths,
            source_filename=filename,
        )
    except ImportError:
        return ParsedDocument(
            text="",
            error="python-docx не установлен на сервере.",
            source_filename=filename,
        )
    except Exception as e:
        return ParsedDocument(text="", error=f"Ошибка DOCX: {e}", source_filename=filename)


async def parse_image(file_bytes: bytes, filename: str, storage, user_id: str) -> ParsedDocument:
    """Use Claude Vision to extract text/description from image."""
    from core.claude import call_claude_vision

    ext = Path(filename).suffix.lstrip(".").lower() or "jpg"
    path = _save_image_binary(storage, file_bytes, ext)

    # Resize if too large for API (base64 limit)
    api_bytes = file_bytes
    if len(file_bytes) > 4 * 1024 * 1024:
        try:
            from PIL import Image as PILImage
            img = PILImage.open(io.BytesIO(file_bytes))
            img.thumbnail((1600, 1600))
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=85)
            api_bytes = buf.getvalue()
        except Exception:
            pass

    b64 = base64.standard_b64encode(api_bytes).decode()
    media_type = _get_media_type(ext)

    try:
        text, _ = await call_claude_vision(
            system=(
                "Ты помощник для извлечения знаний из изображений. "
                "Если есть текст — перепиши его полностью. "
                "Если схема/диаграмма — объясни подробно. "
                "Если фото доски/заметок — перепиши все записи."
            ),
            image_b64=b64,
            media_type=media_type,
            user_id=user_id,
        )
    except Exception as e:
        text = f"[Изображение: {filename}] — не удалось распознать: {e}"

    return ParsedDocument(
        text=text,
        images=[path] if path else [],
        source_filename=filename,
    )


def _get_media_type(ext: str) -> str:
    return {
        "jpg": "image/jpeg", "jpeg": "image/jpeg",
        "png": "image/png", "gif": "image/gif", "webp": "image/webp",
    }.get(ext, "image/jpeg")


async def parse_document(
    file_bytes: bytes,
    filename: str,
    storage,
    user_id: str,
) -> ParsedDocument:
    """Route to correct parser based on file extension."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return await parse_pdf(file_bytes, filename, storage)
    elif ext in (".docx", ".doc"):
        return await parse_docx(file_bytes, filename, storage)
    elif ext in (".jpg", ".jpeg", ".png", ".gif", ".webp"):
        return await parse_image(file_bytes, filename, storage, user_id)
    else:
        try:
            text = file_bytes.decode("utf-8", errors="replace")
            return ParsedDocument(text=text, source_filename=filename)
        except Exception:
            return ParsedDocument(
                text="", error=f"Неизвестный формат: {ext}", source_filename=filename
            )


def format_images_for_obsidian(image_paths: list[str], base_text: str) -> str:
    """Append Obsidian ![[image]] embeds to note body."""
    if not image_paths:
        return base_text
    img_section = "\n\n## Вложения\n"
    for path in image_paths:
        img_section += f"\n![[{Path(path).name}]]"
    return base_text + img_section
